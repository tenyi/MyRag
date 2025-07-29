"""
Word 文件處理器

使用 python-docx 處理 .docx 格式的檔案
"""

import re
from pathlib import Path
from typing import List, Optional

from docx import Document as DocxDocument
from loguru import logger

from ..models.document import Document
from .base import BaseDocumentProcessor
from .exceptions import ContentExtractionError, FileCorruptionError


class DocxProcessor(BaseDocumentProcessor):
    """Word 文件處理器
    
    支援 .docx 格式的檔案
    """
    
    def __init__(self):
        super().__init__()
        self.supported_extensions = {".docx"}
    
    def extract_content(self, file_path: str) -> str:
        """從 Word 檔案中提取文字內容
        
        Args:
            file_path: 檔案路徑
            
        Returns:
            str: 提取的文字內容
            
        Raises:
            ContentExtractionError: 內容提取失敗
            FileCorruptionError: Word 檔案損壞
        """
        try:
            # 開啟 Word 文件
            doc = DocxDocument(file_path)
            
            # 提取段落文字
            paragraphs = []
            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()
                if text:  # 只保留非空段落
                    paragraphs.append(text)
            
            # 提取表格內容
            table_content = self._extract_table_content(doc)
            if table_content:
                paragraphs.extend(table_content)
            
            if not paragraphs:
                raise ContentExtractionError(
                    file_path, "無法從 Word 文件中提取任何內容"
                )
            
            # 合併所有內容
            full_content = "\n\n".join(paragraphs)
            
            # 清理文字內容
            cleaned_content = self._clean_text(full_content)
            
            if not cleaned_content.strip():
                raise ContentExtractionError(
                    file_path, "清理後的內容為空"
                )
            
            return cleaned_content
            
        except Exception as e:
            if isinstance(e, (ContentExtractionError, FileCorruptionError)):
                raise
            
            # 檢查是否為檔案損壞
            if "corrupt" in str(e).lower() or "invalid" in str(e).lower():
                raise FileCorruptionError(
                    file_path, f"Word 檔案可能損壞: {str(e)}"
                )
            
            raise ContentExtractionError(
                file_path, f"Word 文件處理失敗: {str(e)}"
            )
    
    def _extract_table_content(self, doc: DocxDocument) -> List[str]:
        """提取表格內容
        
        Args:
            doc: Word 文件物件
            
        Returns:
            List[str]: 表格內容列表
        """
        table_contents = []
        
        try:
            for table in doc.tables:
                table_text = []
                
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        if cell_text:
                            row_text.append(cell_text)
                    
                    if row_text:
                        table_text.append(" | ".join(row_text))
                
                if table_text:
                    table_contents.append("\n".join(table_text))
            
        except Exception as e:
            logger.warning(f"提取表格內容時發生錯誤: {str(e)}")
        
        return table_contents
    
    def _clean_text(self, text: str) -> str:
        """清理從 Word 文件提取的文字
        
        Args:
            text: 原始文字
            
        Returns:
            str: 清理後的文字
        """
        # 移除多餘的空白字符
        text = re.sub(r'\s+', ' ', text)
        
        # 移除多餘的換行符
        text = re.sub(r'\n\s*\n', '\n\n', text)
        
        # 移除特殊字符（保留中文和基本標點）
        text = re.sub(r'[^\w\s\u4e00-\u9fff，。！？；：「」『』（）【】《》〈〉|]', '', text)
        
        return text.strip()
    
    def get_docx_properties(self, file_path: str) -> dict:
        """取得 Word 文件的屬性
        
        Args:
            file_path: 檔案路徑
            
        Returns:
            dict: 文件屬性
        """
        try:
            doc = DocxDocument(file_path)
            core_props = doc.core_properties
            
            return {
                "title": core_props.title or "",
                "author": core_props.author or "",
                "subject": core_props.subject or "",
                "keywords": core_props.keywords or "",
                "comments": core_props.comments or "",
                "created": core_props.created,
                "modified": core_props.modified,
                "last_modified_by": core_props.last_modified_by or "",
                "revision": core_props.revision,
                "paragraph_count": len(doc.paragraphs),
                "table_count": len(doc.tables),
            }
            
        except Exception as e:
            logger.warning(f"無法讀取 Word 文件屬性: {str(e)}")
            return {}
    
    def process_file(self, file_path: str) -> Document:
        """處理 Word 檔案
        
        Args:
            file_path: 檔案路徑
            
        Returns:
            Document: 文件物件
        """
        # 驗證檔案
        self.validate_file(file_path)
        
        # 提取內容
        content = self.extract_content(file_path)
        
        # 取得文件屬性
        docx_properties = self.get_docx_properties(file_path)
        
        # 使用文件標題作為文件標題（如果有的話）
        title = docx_properties.get("title", "").strip()
        if not title:
            title = Path(file_path).stem
        
        # 建立文件物件
        document = self.create_document(
            file_path=file_path,
            content=content,
            title=title,
            encoding="utf-8"  # Word 內容已轉換為 Unicode
        )
        
        logger.debug(f"成功處理 Word 檔案: {file_path}")
        return document